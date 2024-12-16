"""
File: ui_components.py
Provides all UI components and styling for the AI Hypotheek Assistent.
"""

import streamlit as st
import smtplib
import json
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from io import BytesIO
import logging
import re
from definitions import MORTGAGE_DEFINITIONS, improve_explanation

# Configure logging
logging.basicConfig(filename='app.log', level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def apply_custom_css():
    st.markdown("""
    <style>
    .term-highlight {
        background: linear-gradient(120deg, #E8F0FE 0%, #d2e3fc 100%);
        border-radius: 4px;
        padding: 2px 6px;
        margin: 0 2px;
        cursor: pointer;
        display: inline-block;
        transition: all 0.2s ease;
        box-shadow: 0 1px 2px rgba(0,0,0,0.1);
        color: #1a73e8;
        text-decoration: none;
        font-weight: 500;
    }
    
    .term-highlight:hover {
        background: linear-gradient(120deg, #d2e3fc 0%, #c2d9fc 100%);
        box-shadow: 0 2px 4px rgba(0,0,0,0.15);
        transform: translateY(-1px);
    }
    
    .explanation-card {
        background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
        padding: 1.25rem;
        border-radius: 12px;
        border: 1px solid rgba(226, 232, 240, 0.8);
        margin-bottom: 1rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
    }

    .explanation-title {
        color: #1a73e8;
        margin-bottom: 0.75rem;
        font-size: 1.1rem;
        font-weight: 600;
        border-bottom: 2px solid rgba(26, 115, 232, 0.1);
        padding-bottom: 0.5rem;
    }

    .explanation-content {
        color: #374151;
        font-size: 0.95rem;
        line-height: 1.7;
    }

    .text-paragraph {
        line-height: 1.8;
        margin-bottom: 1.25em;
        color: #1f2937;
        font-size: 1rem;
    }

    .standard-explanations {
        background: #f8fafc;
        padding: 1rem;
        border-radius: 8px;
        margin-top: 1rem;
        border: 1px solid #e2e8f0;
    }

    .qa-history {
        background: #f8fafc;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border: 1px solid #e2e8f0;
    }

    .qa-pair {
        background: white;
        padding: 1rem;
        border-radius: 6px;
        margin: 0.5rem 0;
        border: 1px solid #e5e7eb;
    }

    .qa-context {
        color: #6b7280;
        font-size: 0.9em;
        font-style: italic;
        margin-bottom: 0.5rem;
    }

    .qa-question {
        color: #1a73e8;
        font-weight: 500;
        margin-bottom: 0.5rem;
    }

    .qa-answer {
        color: #374151;
    }

    .loading-overlay {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-color: rgba(255, 255, 255, 0.8);
        display: flex;
        justify-content: center;
        align-items: center;
        z-index: 9999;
    }

    .loading-spinner {
        width: 50px;
        height: 50px;
        border: 5px solid #f3f3f3;
        border-top: 5px solid #3498db;
        border-radius: 50%;
        animation: spin 1s linear infinite;
    }

    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    </style>
    """, unsafe_allow_html=True)

def render_loading_overlay():
    st.markdown("""
        <div class="loading-overlay">
            <div class="loading-spinner"></div>
            <div style="margin-left: 1rem;">Bezig met verwerken...</div>
        </div>
    """, unsafe_allow_html=True)

def render_qa_history(qa_history, section):
    """Renders the Q&A history for a specific section."""
    if not qa_history:
        return
    
    relevant_qa = [qa for qa in qa_history if qa.get('category') == section]
    if not relevant_qa:
        return
    
    st.markdown("""
        <div class="qa-history">
            <h4>Aanvullende informatie uit het gesprek:</h4>
    """, unsafe_allow_html=True)
    
    for qa in relevant_qa:
        st.markdown(f"""
            <div class="qa-pair">
                <div class="qa-context">{qa.get('context', '')}</div>
                <div class="qa-question">Vraag: {qa.get('question', '')}</div>
                <div class="qa-answer">Antwoord: {qa.get('answer', '')}</div>
            </div>
        """, unsafe_allow_html=True)
    
    st.markdown("</div>", unsafe_allow_html=True)

def render_results(app_state):
    """Renders the analysis results with proper validation."""
    st.title("Analyse Resultaten")
    
    # First check if we have a valid transcript
    if not app_state.transcript or not app_state.transcript.strip():
        st.warning("⚠️ Geen transcript beschikbaar. Maak eerst een opname of voer een transcript in.")
        
        # Add a helper button to go back
        if st.button("↩️ Terug naar invoer"):
            app_state.reset()
            st.rerun()
        return
    
    # Then check if we have valid results
    if not app_state.result or not any(content.strip() for content in app_state.result.values()):
        st.warning("⚠️ Geen analyse resultaten beschikbaar.")
        
        # Add a helper button to retry
        if st.button("🔄 Opnieuw analyseren"):
            app_state.reset()
            st.rerun()
        return

    # Show original transcript in expander
    with st.expander("📝 Oorspronkelijk transcript", expanded=False):
        st.markdown(f"```{app_state.transcript}```")

    # Process each section
    for section, content in app_state.result.items():
        # Skip empty sections
        if not content or not content.strip():
            continue
            
        with st.expander(section.replace("_", " ").capitalize(), expanded=True):
            # Display the content
            paragraphs = [p for p in content.split('\n') if p.strip()]
            for para_idx, paragraph in enumerate(paragraphs):
                st.markdown(f'<div class="text-paragraph">{paragraph}</div>', unsafe_allow_html=True)
            
            # Display Q&A history relevant to this section
            render_qa_history(app_state.structured_qa_history, section.replace("adviesmotivatie_", ""))
            
            # Only show standard explanations if we have actual content
            if content.strip():
                render_standard_explanations(section)

    # Export options only if we have valid content
    if any(content.strip() for content in app_state.result.values()):
        st.markdown("### 📑 Exporteer Resultaten")
        if st.button("Exporteer als Word-document", use_container_width=True):
            export_to_docx(app_state)

    if st.button("Nieuwe Analyse", use_container_width=True):
        app_state.reset()
        st.rerun()

def export_to_docx(app_state):
    """Exports the analysis results including Q&A history to a Word document."""
    doc = Document()
    doc.add_heading('Analyse Resultaten', 0)

    # Add original transcript
    doc.add_heading('Oorspronkelijk Transcript', level=1)
    doc.add_paragraph(app_state.transcript)

    # Add analysis results with Q&A history
    for section, content in app_state.result.items():
        doc.add_heading(section.replace("_", " ").capitalize(), level=1)
        
        # Add main content
        for line in content.split('\n'):
            if line.strip().startswith('**') and line.strip().endswith('**'):
                doc.add_paragraph(line.strip()[2:-2], style='Heading 2')
            else:
                doc.add_paragraph(line)
        
        # Add Q&A history for this section
        section_qa = [qa for qa in app_state.structured_qa_history 
                     if qa.get('category') == section.replace("adviesmotivatie_", "")]
        
        if section_qa:
            doc.add_heading('Aanvullende Informatie uit het Gesprek', level=2)
            for qa in section_qa:
                p = doc.add_paragraph()
                p.add_run('Context: ').bold()
                p.add_run(qa.get('context', ''))
                p = doc.add_paragraph()
                p.add_run('Vraag: ').bold()
                p.add_run(qa.get('question', ''))
                p = doc.add_paragraph()
                p.add_run('Antwoord: ').bold()
                p.add_run(qa.get('answer', ''))
                doc.add_paragraph()  # Add spacing

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    st.download_button(
        label="Download Word Document",
        data=bio.getvalue(),
        file_name="analyse_resultaten.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
