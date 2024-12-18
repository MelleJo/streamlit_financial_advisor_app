"""
File: fp_report_service.py
Handles report generation for Financial Planning module.
"""

import logging
from typing import Dict, Any, Optional, BinaryIO
from docx import Document
from docx.shared import Inches, Pt
import io
from datetime import datetime
import json

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FPReportService:
    def __init__(self):
        """Initialize the report service with templates."""
        self.templates = self._load_templates()
        
    def _load_templates(self) -> Dict[str, str]:
        """Load standard text templates for each section."""
        return {
            "nbi": """
            We maken in dit rapport gebruik van de term 'netto besteedbaar inkomen'. 
            Daarmee bedoelen we het geldbedrag dat over blijft wanneer alle lasten zijn betaald.
            
            In uw geval hebben we uitgegaan van een gewenst netto besteedbaar inkomen van € {nbi} per jaar.
            Dit bedrag is gebaseerd op {toelichting}.
            """,
            
            "pensioen_voldoende": """
            Wanneer u met pensioen gaat beschikt u over voldoende inkomen. 
            Het verwachte netto besteedbaar inkomen bedraagt circa € {pensioen_inkomen} per jaar.
            
            Dit is voldoende om uw gewenste levensstandaard van € {gewenst_inkomen} per jaar te handhaven.
            """,
            
            "pensioen_onvoldoende": """
            Uit de analyse blijkt dat er een tekort ontstaat bij pensionering.
            Het verwachte netto besteedbaar inkomen van € {pensioen_inkomen} per jaar
            is onvoldoende voor uw gewenste levensstandaard van € {gewenst_inkomen} per jaar.
            
            Wij adviseren u om aanvullende maatregelen te treffen, zoals {maatregelen}.
            """,
            
            "ao_zelfstandige": """
            Als zelfstandig ondernemer is het van belang om goed verzekerd te zijn tegen arbeidsongeschiktheid.
            
            {if_aov}Bij arbeidsongeschiktheid ontvangt u een uitkering van € {uitkering} per jaar
            van uw AOV bij {verzekeraar}.
            
            {if_no_aov}U heeft momenteel geen arbeidsongeschiktheidsverzekering.
            Wij adviseren u om een AOV af te sluiten met de volgende dekking:
            - Verzekerd bedrag: € {advies_bedrag}
            - Eigen risicoperiode: {eigen_risico}
            - Eindleeftijd: {eindleeftijd}
            """,
            
            "ao_loondienst": """
            Bij arbeidsongeschiktheid heeft u vanuit uw werkgever en het UWV de volgende dekkingen:
            {dekkingen}
            
            {if_wlv}Aanvullend heeft u een woonlastenverzekering met een dekking van € {wlv_dekking}.
            
            {if_no_wlv}Wij adviseren u een woonlastenverzekering af te sluiten met een dekking
            van € {advies_dekking} om het inkomensverlies bij arbeidsongeschiktheid op te vangen.
            """,
            
            "overlijden": """
            We hebben het overlijdensscenario geanalyseerd voor uw situatie.
            
            {if_voldoende}De huidige voorzieningen zijn voldoende om de gewenste levensstandaard
            te handhaven bij overlijden.
            
            {if_onvoldoende}Er ontstaat een tekort bij overlijden. Wij adviseren een 
            overlijdensrisicoverzekering met een dekking van € {orv_dekking} af te sluiten
            bij {verzekeraar}. De maandpremie bedraagt € {orv_premie}.
            """,
            
            "werkloosheid": """
            Bij werkloosheid heeft u recht op een WW-uitkering voor maximaal {ww_duur} maanden.
            
            {if_voldoende}Samen met uw buffer van € {buffer} is dit voldoende om de periode
            van werkloosheid te overbruggen.
            
            {if_onvoldoende}Wij adviseren u om uw buffer van € {buffer} te verhogen naar minimaal
            € {gewenste_buffer} om een periode van werkloosheid te kunnen overbruggen.
            """
        }

    def generate_report(self, fp_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate the complete FP report by combining analysis with templates."""
        try:
            report = {}
            
            # Generate each section using templates
            report["samenvatting"] = self._generate_summary(fp_data)
            report["huidige_situatie"] = self._generate_current_situation(fp_data)
            report["pensioen"] = self._generate_pension_section(fp_data)
            report["risico_analyse"] = self._generate_risk_analysis(fp_data)
            report["vermogen"] = self._generate_wealth_section(fp_data)
            report["actiepunten"] = self._generate_action_points(fp_data)
            
            # Add graph placeholders
            report["graphs"] = self._generate_graph_placeholders(fp_data)
            
            return report
            
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}")
            return self._get_default_report()

    def _generate_summary(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate the summary section."""
        try:
            nbi_template = self.templates["nbi"]
            summary = nbi_template.format(
                nbi=data.get("netto_besteedbaar_inkomen", "0"),
                toelichting=data.get("nbi_toelichting", "uw opgegeven wensen en doelstellingen")
            )
            
            return {
                "content": summary,
                "nbi": data.get("netto_besteedbaar_inkomen"),
                "doelstellingen": data.get("doelstellingen", []),
                "kernadvies": data.get("kernadvies", "")
            }
            
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}")
            return {}

    def _generate_pension_section(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate the pension section using appropriate template."""
        try:
            pension_data = data.get("pensioen", {})
            is_sufficient = pension_data.get("is_sufficient", False)
            
            template = self.templates["pensioen_voldoende" if is_sufficient else "pensioen_onvoldoende"]
            content = template.format(
                pensioen_inkomen=pension_data.get("verwacht_inkomen", "0"),
                gewenst_inkomen=pension_data.get("gewenst_inkomen", "0"),
                maatregelen=pension_data.get("maatregelen", "aanvullende pensioenopbouw")
            )
            
            return {
                "content": content,
                "graph_data": self._prepare_pension_graph_data(pension_data),
                "is_sufficient": is_sufficient,
                "details": pension_data
            }
            
        except Exception as e:
            logger.error(f"Error generating pension section: {str(e)}")
            return {}

    def _generate_risk_analysis(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate the risk analysis section."""
        try:
            risk_data = data.get("risicos", {})
            
            # Generate AO (Arbeidsongeschiktheid) section
            ao_content = self._generate_ao_section(risk_data.get("ao", {}))
            
            # Generate overlijden section
            overlijden_content = self._generate_overlijden_section(risk_data.get("overlijden", {}))
            
            # Generate werkloosheid section
            werkloosheid_content = self._generate_werkloosheid_section(risk_data.get("werkloosheid", {}))
            
            return {
                "ao": ao_content,
                "overlijden": overlijden_content,
                "werkloosheid": werkloosheid_content,
                "graphs": self._prepare_risk_graphs(risk_data)
            }
            
        except Exception as e:
            logger.error(f"Error generating risk analysis: {str(e)}")
            return {}

    def _generate_ao_section(self, ao_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate arbeidsongeschiktheid section based on employment status."""
        try:
            is_zelfstandig = ao_data.get("is_zelfstandig", False)
            template = self.templates["ao_zelfstandige" if is_zelfstandig else "ao_loondienst"]
            
            if is_zelfstandig:
                has_aov = ao_data.get("has_aov", False)
                content = template.format(
                    if_aov="" if not has_aov else "# ",
                    if_no_aov="" if has_aov else "# ",
                    uitkering=ao_data.get("aov_uitkering", "0"),
                    verzekeraar=ao_data.get("aov_verzekeraar", ""),
                    advies_bedrag=ao_data.get("advies_bedrag", "0"),
                    eigen_risico=ao_data.get("advies_eigen_risico", "90 dagen"),
                    eindleeftijd=ao_data.get("advies_eindleeftijd", "67")
                )
            else:
                has_wlv = ao_data.get("has_wlv", False)
                content = template.format(
                    dekkingen=ao_data.get("dekkingen", ""),
                    if_wlv="" if not has_wlv else "# ",
                    if_no_wlv="" if has_wlv else "# ",
                    wlv_dekking=ao_data.get("wlv_dekking", "0"),
                    advies_dekking=ao_data.get("advies_dekking", "0")
                )
            
            return {
                "content": content,
                "details": ao_data
            }
            
        except Exception as e:
            logger.error(f"Error generating AO section: {str(e)}")
            return {}

    def _generate_overlijden_section(self, overlijden_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate overlijden section."""
        try:
            is_sufficient = overlijden_data.get("is_sufficient", False)
            template = self.templates["overlijden"]
            
            content = template.format(
                if_voldoende="" if not is_sufficient else "# ",
                if_onvoldoende="" if is_sufficient else "# ",
                orv_dekking=overlijden_data.get("orv_dekking", "0"),
                verzekeraar=overlijden_data.get("verzekeraar", ""),
                orv_premie=overlijden_data.get("orv_premie", "0")
            )
            
            return {
                "content": content,
                "details": overlijden_data
            }
            
        except Exception as e:
            logger.error(f"Error generating overlijden section: {str(e)}")
            return {}

    def _generate_werkloosheid_section(self, werkloosheid_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate werkloosheid section."""
        try:
            is_sufficient = werkloosheid_data.get("is_sufficient", False)
            template = self.templates["werkloosheid"]
            
            content = template.format(
                ww_duur=werkloosheid_data.get("ww_duur", "0"),
                if_voldoende="" if not is_sufficient else "# ",
                if_onvoldoende="" if is_sufficient else "# ",
                buffer=werkloosheid_data.get("buffer", "0"),
                gewenste_buffer=werkloosheid_data.get("gewenste_buffer", "0")
            )
            
            return {
                "content": content,
                "details": werkloosheid_data
            }
            
        except Exception as e:
            logger.error(f"Error generating werkloosheid section: {str(e)}")
            return {}

    def _prepare_pension_graph_data(self, pension_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare data for pension graphs."""
        return {
            "inkomen_prognose": {
                "title": "Pensioeninkomen Prognose",
                "type": "line",
                "data_requirements": {
                    "x_axis": "Leeftijd",
                    "y_axis": "Inkomen (€)",
                    "series": ["Verwacht Inkomen", "Gewenst Inkomen"]
                }
            },
            "vermogen_prognose": {
                "title": "Pensioenvermogen Prognose",
                "type": "line",
                "data_requirements": {
                    "x_axis": "Jaar",
                    "y_axis": "Vermogen (€)",
                    "series": ["Opgebouwd Vermogen", "Doelvermogen"]
                }
            }
        }

    def _prepare_risk_graphs(self, risk_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare data for risk analysis graphs."""
        return {
            "ao_inkomen": {
                "title": "Inkomen bij Arbeidsongeschiktheid",
                "type": "line",
                "data_requirements": {
                    "x_axis": "Jaar",
                    "y_axis": "Inkomen (€)",
                    "series": ["Inkomen bij AO", "Benodigd Inkomen"]
                }
            },
            "overlijden_inkomen": {
                "title": "Inkomen bij Overlijden",
                "type": "line",
                "data_requirements": {
                    "x_axis": "Jaar",
                    "y_axis": "Inkomen (€)",
                    "series": ["Nabestaandeninkomen", "Benodigd Inkomen"]
                }
            }
        }

    def generate_word_report(self, fp_data: Dict[str, Any]) -> BinaryIO:
        """Generate a Word document version of the report."""
        doc = Document()
        
        # Add title
        doc.add_heading('Financieel Plan', 0)
        doc.add_paragraph(f'Gegenereerd op: {datetime.now().strftime("%d-%m-%Y")}')
        
        # Add sections
        self._add_summary_section(doc, fp_data.get("samenvatting", {}))
        self._add_pension_section(doc, fp_data.get("pensioen", {}))
        self._add_risk_section(doc, fp_data.get("risico_analyse", {}))
        self._add_action_points(doc, fp_data.get("actiepunten", {}))
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_summary_section(self, doc: Document, summary_data: Dict[str, Any]):
        """Add summary section to Word document."""
        doc.add_heading('Samenvatting', level=1)
        doc.add_paragraph(summary_data.get("content", ""))
        
        if summary_data.get("doelstellingen"):
            doc.add_heading('Doelstellingen', level=2)
            for doel in summary_data["doelstellingen"]:
                doc.add_paragraph(doel, style='List Bullet')

    def _add_pension_section(self, doc: Document, pension_data: Dict[str, Any]):
        """Add pension section to Word document."""
        doc.add_heading('Pensioen', level=1)
        doc.add_paragraph(pension_data.get("content", ""))
        
        # Add graph placeholders
        if pension_data.get("graphs"):
            doc.add_paragraph("Hier worden de volgende grafieken ingevoegd:")
            for graph in pension_data["graphs"].values():
                doc.add_paragraph(f"- {graph['title']}", style='List Bullet')

    def _add_risk_section(self, doc: Document, risk_data: Dict[str, Any]):
        """Add risk analysis section to Word document."""
        doc.add_heading('Risico Analyse', level=1)
        
        # Add subsections
        for risk_type in ['ao', 'overlijden', 'werkloosheid']:
            if risk_data.get(risk_type):
                doc.add_heading(risk_type.capitalize(), level=2)
                doc.add_paragraph(risk_data[risk_type].get("content", ""))

    def _add_action_points(self, doc: Document, action_points: Dict[str, Any]):
        """Add action points to Word document."""
        doc.add_heading('Actiepunten', level=1)
        
        if action_points.get("client"):
            doc.add_heading('Voor Client', level=2)
            for action in action_points["client"]:
                doc.add_paragraph(action, style='List Bullet')
                
        if action_points.get("veldhuis"):
            doc.add_heading('Voor Veldhuis', level=2)
            for action in action_points["veldhuis"]:
                doc.add_paragraph(action, style='List Bullet')

    def _get_default_report(self) -> Dict[str, Any]:
        """Return default report structure when generation fails."""
        return {
            "samenvatting": {"content": "Rapport generatie niet mogelijk"},
            "pensioen": {"content": ""},
            "risico_analyse": {},
            "actiepunten": {"client": [], "veldhuis": []},
            "graphs": {}
        }